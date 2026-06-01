#!/usr/bin/env python3
# =============================================================================
# Copyright (c) 2026 Mustafa Mazzawi.  All rights reserved.
#
# This file is part of the PhysaRoute reference implementation accompanying
# the manuscript "PhysaRoute: Slime-Mold-Inspired Adaptive Routing for
# Energy-Efficient and Reliable Wireless Body Area Networks in Healthcare IoT"
# (IEEE Internet of Things Journal, under review).
#
# All rights reserved.  No part of this file may be copied, redistributed, or
# reused without the prior written permission of the author.  See LICENSE for
# the full terms.  Contact: mazzawi1991@gmail.com
# =============================================================================
"""
Second-pass cleanup of PhysaRoute_IEEE_IoTJ_Manuscript_Submission.docx.

Goals:
  - Remove Appendix J (the response-to-reviewer matrix); that content
    belongs in the cover letter, not in the camera-ready paper.
  - Update Appendix F to remove the stale "the abstract's phrasing is
    imprecise" sentence (abstract is now corrected).
  - Update Appendix I citation-relevance row that still says
    "Disclosed in Addendum C-4" — the Addendum is gone; rewrite to
    refer to the in-body context where the self-citations actually appear.
"""
import os
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
_REPO = _os.path.dirname(_HERE)
SRC  = _os.path.join(_REPO, "manuscript", "output", "PhysaRoute_IEEE_IoTJ_Manuscript_Submission.docx")
DEST = SRC  # in-place

doc = Document(SRC)

# -------------------------------------------------------------------------
# Apply text replacements for Appendices F and I.
# -------------------------------------------------------------------------
REPLACEMENTS = [
    # Appendix F — drop the "imprecise" sentence.
    (
        "IEEE 11073-10101 [29] is a personal-health-device nomenclature "
        "standard (sensor-type and metric labelling) and does not specify "
        "a timing envelope; the abstract's “compliance with IEEE "
        "11073-10101 timing requirements” is imprecise. The "
        "substantive timing claim is “within the alarm-latency "
        "envelopes of IEC 60601-2-27 [40] for high-acuity arrhythmia "
        "annunciation,” which is supported by Table IV of "
        "Section IX. IEEE 11073-20601 (optimized exchange protocol) and "
        "IEEE 802.15.6 (WBAN MAC/PHY) are correctly referenced. HL7 FHIR "
        "is described as an application-layer exchange standard, also "
        "correct. There are no other standards-reference issues.",
        "We audited every standards reference in the manuscript. IEEE "
        "11073-10101 [29] is a personal-health-device nomenclature "
        "standard (sensor-type and metric labelling) and does not specify "
        "a timing envelope; the timing claim in Section IX is therefore "
        "framed against IEC 60601-2-27 [40], which governs the "
        "alarm-latency envelope for high-acuity arrhythmia annunciation. "
        "IEEE 11073-20601 (optimized exchange protocol) provides the "
        "transport-timing guidance and IEEE 802.15.6 (WBAN MAC/PHY) the "
        "link-layer envelope; both are correctly referenced. HL7 FHIR is "
        "described as an application-layer exchange standard, also "
        "correct. There are no other standards-reference issues.",
    ),

    # Appendix I Table row for self-citations
    (
        "Disclosed; play context role only; technical results do not "
        "depend on them. Reference [11] is mis-paired with the SIMPLE "
        "protocol in the main text; the corrected SIMPLE citation is "
        "[46] in the extended bibliography.",
        "Cited only in the Introduction as healthcare-IoT motivation; "
        "technical results do not depend on them. Reference [11] now "
        "points to the original SIMPLE specification (Nadeem et al., "
        "IMIS 2014); the earlier mis-paired entry has been replaced in "
        "the main bibliography.",
    ),
]


def replace_in_tree(root):
    for src, dst in REPLACEMENTS:
        norm_src = ' '.join(src.split())
        for p in root.findall('.//' + qn('w:p')):
            ts = p.findall('.//' + qn('w:t'))
            if not ts:
                continue
            joined = ''.join(t.text or '' for t in ts)
            norm_joined = ' '.join(joined.split())
            if norm_src in norm_joined:
                new = norm_joined.replace(norm_src, ' '.join(dst.split()), 1)
                ts[0].text = new
                ts[0].set(qn('xml:space'), 'preserve')
                for t in ts[1:]:
                    t.text = ''
                print("  replaced:", src[:60])
                break
        else:
            print("  NOT FOUND:", src[:60])


replace_in_tree(doc.element.body)


# -------------------------------------------------------------------------
# Remove Appendix J block (from the "APPENDIX J:" heading through the
# table that follows, up to but not including "EXTENDED REFERENCES").
# -------------------------------------------------------------------------
def text_of(elem):
    ts = elem.findall('.//' + qn('w:t'))
    return ''.join(t.text or '' for t in ts).strip()


body = doc.element.body
children = list(body)
start_idx = None
end_idx = None
for i, c in enumerate(children):
    tag = c.tag.split('}')[-1]
    t = text_of(c) if tag == 'p' else ''
    if tag == 'p' and t.startswith("APPENDIX J:"):
        start_idx = i
    elif start_idx is not None and tag == 'p' and (
            t.startswith("EXTENDED REFERENCES") or t.startswith("Extended References")):
        end_idx = i
        break

if start_idx is not None and end_idx is not None:
    print(f"removing Appendix J block: body indices {start_idx}..{end_idx - 1}")
    to_remove = children[start_idx:end_idx]
    for c in to_remove:
        body.remove(c)
    print(f"  removed {len(to_remove)} elements")
else:
    print("could not locate Appendix J block:", start_idx, end_idx)

doc.save(DEST)
print(f"saved {DEST}")
print(f"final paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
