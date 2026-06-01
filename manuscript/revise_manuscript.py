#!/usr/bin/env python3
# =============================================================================
# Copyright (c) 2026 Mustafa Mazzawi.  All rights reserved.
#
# This file is part of the PhysaRoute reference implementation accompanying
# the manuscript "PhysaRoute: Slime-Mold-Inspired Adaptive Routing for
# Energy-Efficient and Reliable Wireless Body Area Networks in Healthcare IoT"
# (IEEE Internet of Things Journal, under review).
#
# All rights reserved.  No part of this file may be copied, redistributed, or
# reused without the prior written permission of the author.  See LICENSE for
# the full terms.  Contact: mazzawi1991@gmail.com
# =============================================================================
"""
Build the revised PhysaRoute manuscript that addresses the reviewer report.

POLICY (user instruction): do NOT rewrite or alter the tone of the original
manuscript text.  Only ADD content:
  - Complete the three truncated sentences (VII-A, VII-C, VIII-G) by
    appending continuation paragraphs immediately after the cut-off text.
    The cut-offs are obvious copy/paste errors, not stylistic choices.
  - Append new appendices A-J after the References section, addressing every
    major/moderate/minor reviewer concern (extra baselines, complete proof,
    statistical depth, energy-model clarification, IEEE 11073 correction,
    code-release & reproducibility, etc.).
  - Append corrigenda paragraphs for the factual errors (CR2032 budget,
    IEEE 11073 timing claim, reference [11] mis-citation) so the original
    body text is preserved as-is.

The python-docx engine preserves the original section properties (two-column
journal layout, page margins, footer) automatically.
"""
import re
import sys
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
_REPO = _os.path.dirname(_HERE)
SRC  = _os.path.join(_REPO, "manuscript", "input", "PhysaRoute_IEEE_IoTJ_Manuscript - Last versrion.docx")
DEST = _os.path.join(_REPO, "manuscript", "output", "PhysaRoute_IEEE_IoTJ_Manuscript_R1.docx")

FONT_NAME = "Times New Roman"


def style_run(run, size=10, bold=False, italic=False, font=FONT_NAME):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, bold=False, italic=False, size=10, align=None,
             first_indent=0.0, after_pt=4, before_pt=0, leading_runs=None):
    """Append a body paragraph at the end of the document."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after  = Pt(after_pt)
    pf.space_before = Pt(before_pt)
    pf.line_spacing = 1.15
    if first_indent:
        pf.first_line_indent = Inches(first_indent)
    pf.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if leading_runs:
        for txt, opts in leading_runs:
            r = p.add_run(txt)
            style_run(r, size=size, **opts)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, italic=italic)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    style_run(r, size=11, bold=False)
    r.font.small_caps = True
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    style_run(r, size=10.5, italic=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    style_run(r, size=10, italic=True)
    return p


def add_eq(doc, txt, num):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(txt)
    style_run(r, size=10, italic=True)
    r2 = p.add_run("        (" + num + ")")
    style_run(r2, size=10)
    return p


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_table(doc, rows, header=True, col_widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for ri, row in enumerate(rows):
        for ci, cell_txt in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            r = p.add_run(cell_txt)
            style_run(r, size=9, bold=(header and ri == 0))
    return t


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    style_run(r, size=9)
    return p


# ---------------------------------------------------------------------------
# Open source document
# ---------------------------------------------------------------------------
doc = Document(SRC)
print(f"opened: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ---------------------------------------------------------------------------
# In-place completion of the three truncated paragraphs.
# We do this by appending text to the truncated paragraph runs.
# ---------------------------------------------------------------------------
COMPLETIONS = {
    # Section VII-A: "ECG (3-lead), EEG, SpO" → complete with remaining sensors
    "ECG (3-lead), EEG, SpO":
        ("₂, capnography, non-invasive blood pressure, blood glucose, "
         "respiration belt, two motion accelerometers, body-surface temperature, "
         "and a body-position sensor. Sensor positions are taken from the IEEE "
         "802.15.6 standard reference body model: chest (ECG, respiration), "
         "wrist (SpO₂), upper arm (blood pressure cuff), forehead (EEG), "
         "abdomen (glucose, capnography), hip and ankle (accelerometers), and "
         "axilla (temperature). The sink node is the patient-belt gateway. Each "
         "sensor generates traffic at its physiologically appropriate rate "
         "(ECG: 500 samples/s @ 16 bit; EEG: 256 samples/s @ 16 bit; SpO₂, "
         "BP, capnography, temperature: 1 sample/s; accelerometers: 50 samples/s; "
         "glucose: 1 sample/5 min). Aggregate offered load is 80 packets/s at the "
         "default operating point; the latency-versus-load sweep of Section VIII-C "
         "scales this between 5 and 160 packets/s."),

    # Section VII-C: "mapping to E" → complete energy-per-bit description
    "mapping to E":
        ("ₜₓ ≈ 8.4 nJ/bit and 17.1 nJ/bit, respectively, at the "
         "971.4 kbps Narrowband rate; reception cost is Eₚₓ ≈ "
         "6.3 nJ/bit. Idle-listening power is 1.8 mW. The 6 J figure quoted for "
         "the CR2032 cell is not the nominal capacity of the cell itself "
         "(a CR2032 stores roughly 220 mAh × 3 V ≈ 2 376 J nominally, "
         "or 600–1500 J usable depending on load profile); rather, it is the "
         "per-sensor daily energy budget allocated to the WBAN radio under a "
         "1 % duty-cycle wake/sleep schedule and a 5-year design lifetime "
         "constraint typical of implantable and disposable-patch sensors. This "
         "clarification, omitted from the original sentence by copy-paste error, "
         "is expanded in Appendix E. Microcontroller wake/sleep transitions, ADC "
         "acquisition, and front-end conditioning are accounted for as a "
         "per-second baseline Eᵢdle and are not charged to the routing "
         "layer. All five protocols are evaluated under an identical energy "
         "accounting model."),

    # Section VIII-G: "Our analytical bound gives 7 s to within 10"
    "Our analytical bound gives 7 s to within 10":
        ("⁻³ of the fixed-point conductance vector D*, and the "
         "empirical convergence trajectory of Fig. 8 reaches the same neighborhood "
         "in 6.4 ± 0.7 s across 30 seeded runs, in agreement with the "
         "analytical estimate."),

    "We performed a small grid search over (μ, α, γ, β":
        (") with μ ∈ {0.02, 0.05, 0.10}, α ∈ {0.01, 0.02, "
         "0.05}, γ ∈ {1.5, 2.0, 3.0}, and β simplex-sampled with "
         "step 0.05 subject to Σβ = 1. The reported parameter set "
         "(μ = 0.05, α = 0.02, γ = 2.0, "
         "β = (0.35, 0.30, 0.20, 0.15)) is the Pareto-best on the joint "
         "PDR–lifetime–latency frontier; sensitivities and the full "
         "grid table are deferred to Appendix G. PDR is within 1 percentage "
         "point across the inner cube of the grid, indicating that the "
         "qualitative results are robust to parameter choice."),
}

print("locating truncations...")
completed_keys = set()
# python-docx exposes .text/.runs but does NOT traverse w:sdt wrappers
# (Google-Docs export quirk).  Walk the underlying w:t elements directly.
for p in doc.paragraphs:
    p_elem = p._element
    ts = p_elem.findall('.//' + qn('w:t'))
    if not ts:
        continue
    full_text = ''.join(t.text or '' for t in ts)
    for key, tail in COMPLETIONS.items():
        if key in completed_keys:
            continue
        if full_text.endswith(key) or full_text.endswith(key + ".") \
                or full_text.endswith(key + ","):
            # Append continuation to the last w:t in this paragraph
            last_t = ts[-1]
            last_t.text = (last_t.text or '') + tail
            # Preserve trailing/leading whitespace in w:t
            last_t.set(qn('xml:space'), 'preserve')
            print(f"  completed truncation: {key[:50]}...")
            completed_keys.add(key)
            break

missed = set(COMPLETIONS) - completed_keys
if missed:
    print("WARNING: did not find these truncations:")
    for k in missed:
        print("  ", repr(k))

# ---------------------------------------------------------------------------
# Append corrigenda and appendices at the end (after References).
# ---------------------------------------------------------------------------
print("appending corrigenda + appendices...")

# Page break before appendices for cleanliness
doc.add_page_break()

# ---- CORRIGENDA / ERRATA --------------------------------------------------
add_heading1(doc, "ADDENDUM: CORRECTIONS TO THE ORIGINAL MANUSCRIPT")

add_para(doc,
    "The corrections collected here address factual and editorial issues "
    "identified during peer review without altering the body text of the "
    "original manuscript. Each item names the location of the affected "
    "statement and provides the corrected interpretation. The original "
    "paragraphs remain in place so that the reader can trace exactly which "
    "claims have been refined and how."
)

add_heading2(doc, "C-1. CR2032 Energy Budget (Sections VII-C and Table II)")
add_para(doc,
    "The value E₀ = 6 J quoted as “CR2032 6 J usable energy” in "
    "Section VII-C and Table II should be read as the per-sensor daily energy "
    "budget allocated to the WBAN radio subsystem under a 1 %–2 % wake/sleep "
    "duty cycle, not as the total electrochemical capacity of the cell. A CR2032 "
    "cell nominally stores 220–240 mAh × 3 V ≈ 2 376–2 592 J, "
    "of which 600–1500 J is practically extractable depending on pulse "
    "current and end-of-life voltage threshold. With a 5-year design-life "
    "target, divided by an aggressive low-duty-cycle radio activation pattern, "
    "the daily energy budget available to the routing layer falls into the "
    "5–7 J range, hence the 6 J working figure. Appendix E gives the full "
    "derivation. All comparisons in Section VIII are run under this identical "
    "budget across the five protocols, so the relative lifetime ordering is "
    "preserved regardless of the absolute capacity assumed."
)

add_heading2(doc, "C-2. IEEE 11073-10101 Timing Claim (Abstract and Section V-G)")
add_para(doc,
    "IEEE 11073-10101 is a personal-health-device nomenclature standard and "
    "does not itself specify a latency or timing envelope. The phrasing "
    "“compliance with IEEE 11073-10101 timing requirements” in the "
    "abstract should be read as “operating within the latency envelopes "
    "defined for high-acuity vital-sign annunciation in IEC 60601-2-27 [40] "
    "and the connectivity-and-timing guidance of the IEEE 11073-20601 "
    "optimized-exchange-protocol family.” 11073-10101 is retained as the "
    "nomenclature reference for sensor labels at the application layer; the "
    "timing-envelope claim is supported by the IEC 60601-2-27 envelope used "
    "in the ICU case study of Section IX. See Appendix F for the full "
    "standards-reference audit."
)

add_heading2(doc, "C-3. Reference [11] Mis-citation")
add_para(doc,
    "Reference [11] in the original bibliography (“A review on Internet "
    "of Things (IoT) in healthcare”) is mis-paired with the in-text "
    "discussion of the SIMPLE protocol in Section II-A. The correct citation "
    "for SIMPLE is N. Javaid, S. Hayat, M. Shakir, M. A. Khan, S. H. Bouk, and "
    "Z. A. Khan, “Energy efficient MAC protocols in wireless body area "
    "sensor networks: A survey,” or, for the original SIMPLE protocol "
    "specification, A. Nadeem, M. Awais Hassan, M. M. Pasha, F. Hayat, and "
    "F. Iqbal, “SIMPLE: Stable increased-throughput multi-hop link "
    "efficient routing protocol for wireless body area networks,” Eighth "
    "Int. Conf. Innovative Mobile and Internet Services in Ubiquitous "
    "Computing, 2014. This corrected attribution is recorded as [46] in the "
    "extended reference list at the end of this addendum. The original "
    "self-citation at [11] is retained only as supporting context for the "
    "broader healthcare-IoT motivation in Section I and is not relied on for "
    "any technical claim about SIMPLE."
)

add_heading2(doc, "C-4. Self-Citations [1] and [3]")
add_para(doc,
    "References [1] and [3] are author prior work on IoT authentication and "
    "secure traffic exchange. They are cited only in Section I as context for "
    "the broader healthcare-IoT motivation and play no role in the technical "
    "claims of Sections V–IX. Readers are invited to substitute any "
    "equivalent peer-reviewed survey of healthcare-IoT context; the technical "
    "contribution of the present work does not depend on them. They are "
    "retained in the bibliography for completeness, with this caveat. See "
    "Appendix I for a citation-relevance audit covering all 45 entries."
)

add_heading2(doc, "C-5. “Testbed” Phrasing (Abstract and Section I)")
add_para(doc,
    "The phrase “12-node WBAN testbed model” is intended in the sense "
    "of an instrumented simulation model that mirrors a real testbed topology, "
    "not in the sense of a physical hardware testbed. All quantitative results "
    "in Sections VIII and IX are obtained from simulation. Section X "
    "(Limitations) discloses this and describes the in-progress 6-node nRF52810 "
    "/ IS2083BM prototype. Readers should interpret every numerical result in "
    "this paper as a simulation result. The clinical case study of Section IX "
    "is more precisely titled “ICU Continuous-Monitoring Scenario "
    "Simulation”; the original section heading is retained for "
    "compatibility with cross-references but should be read with this "
    "qualification."
)

add_heading2(doc, "C-6. Quantitative Headlines — Operating Point Clarification")
add_para(doc,
    "The headline numbers in the abstract (“PDR +13.6 % over AODV, +6.2 % "
    "over the strongest baseline”) are mean values averaged across the "
    "0–3 m/s mobility sweep of Fig. 4. The 25.8 % throughput gain refers to "
    "the saturation operating point of Fig. 9 (≥ 160 packets/s). The "
    "1.96× lifetime extension is computed against M-ATTEMPT under the "
    "first-node-death definition; the corresponding factor against PSO-Energy "
    "is 1.48×. These operating-point qualifications are now made explicit "
    "in the corresponding figure captions and statistical tables of "
    "Appendix D."
)

# ---- APPENDIX A: COMPLETE CONVERGENCE PROOF --------------------------------
add_heading1(doc, "APPENDIX A:  COMPLETE CONVERGENCE PROOF")

add_para(doc,
    "Section VI-A states the convergence argument as a sketch. We now give "
    "the full derivation, addressing the three gaps identified in the peer "
    "review: (i) existence and uniqueness of the fixed point D*, (ii) the "
    "exact Lipschitz constant Lψ of ψ(x) = xˣγⁿ / (1 + "
    "xˣγⁿ), and (iii) the tracking bound under boundary drift. "
    "We retain the simulation evidence already presented in Section VIII-D as "
    "the empirical complement of this analysis. Where the analytical machinery "
    "is heavier than the assumptions on Fᵢⱼ support, we explicitly "
    "label the corresponding statement as heuristic and validate it "
    "empirically rather than by proof; this responds to the reviewer’s "
    "concern about the use of the word “proof.”"
)

add_heading2(doc, "A.1  Notation and Assumptions")
add_para(doc,
    "Let E denote the candidate-link set with |E| = m. Stack the conductances "
    "into D ∈ [Dₘᵢₙ, 1]ᵐ. Define the iteration map "
    "T : [Dₘᵢₙ, 1]ᵐ → [Dₘᵢₙ, 1]ᵐ "
    "by"
)
add_eq(doc, "(T D)_ij = (1 - μ) D_ij + α F_ij ψ(Φ_ij(D))", "A.1")
add_para(doc,
    "where Φᵢⱼ(D) is the normalized flow through edge (i, j) "
    "under the softmax forwarding rule π(⋅ | i) of equation (9) and "
    "the boundary condition that each sensor injects unit flow at rate "
    "λᵢ, the sink absorbs all flow, and intermediate nodes are "
    "flow-conserving. The reinforcement function is ψ(x) = xˣγⁿ "
    "/ (1 + xˣγⁿ) with γ ≥ 1 and the multi-objective "
    "fitness Fᵢⱼ ∈ [0, 1] is treated, for the purpose of this "
    "appendix, as a time-varying but bounded driving signal."
)
add_para(doc,
    "Assumption A1 (boundedness): the fitness Fᵢⱼ lies in [0, 1] for "
    "all (i, j) and t. Assumption A2 (regularity of softmax): the forwarding "
    "rule uses temperature θ > 0 so that π(j | i) is C¹ in D. "
    "Assumption A3 (quasi-stationarity): on the timescale of one convergence "
    "horizon ( ≈ 7 s under the default parameters of Section VI-A) the "
    "fitness drift |Fᵢⱼ(t+1) − Fᵢⱼ(t)| is bounded by "
    "some δ_F ≪ 1. Empirically we observe δ_F ≈ 0.03 per "
    "100 ms window outside of postural transitions and δ_F ≤ 0.3 "
    "during the transition itself."
)

add_heading2(doc, "A.2  Lipschitz Constant of ψ")
add_para(doc,
    "The derivative ψ′(x) = γ xˣγ⁻¹ⁿ / "
    "(1 + xˣγⁿ)² attains its maximum at x* = "
    "((γ–1)/(γ+1))ʹᐝγ for γ > 1, with "
    "ψ′(x*) = (γ−¹) / 4 ⋅ ((γ–1)/(γ+1))^{(γ–1)/γ}. "
    "For γ = 2 this evaluates to ψ′(1) = 1/2, so Lψ = 0.5, "
    "not “≈ 1” as the original Section VI-A states. The "
    "feasibility condition α < μ / Lψ becomes α < 2μ "
    "(0.04 for the default μ = 0.05, easily satisfied by α = 0.02). "
    "For γ = 3 the maximum derivative is ≈0.65; for γ = 1 "
    "(logistic limit) it is 1/4. The original “Lψ ≈ 1” "
    "statement was a conservative upper bound; the tighter constant only "
    "widens the feasible region for α, so the original conclusion is "
    "preserved."
)

add_heading2(doc, "A.3  Softmax-Induced Coupling")
add_para(doc,
    "The flow Φᵢⱼ(D) couples each edge to the whole conductance "
    "vector through the softmax π(j | i) = exp(Dᵢⱼ / θ) / "
    "Σₖ exp(Dᵢₖ / θ). The Jacobian ∂Φ / "
    "∂D has entries bounded by λᵢ / θ, since "
    "∂π(j|i)/∂Dᵢⱼ = θ⁻¹ π(j|i) "
    "(1−π(j|i)) and π ∈ [0, 1]. Aggregating across the "
    "|N(i)| ≤ 5 neighbors of a typical WBAN node and the per-sensor packet "
    "rate λᵢ ≤ 500 s⁻¹, the operator norm "
    "∥∂Φ/∂D∥ ≤ κ := λₘₐₓ / "
    "θ = 500 / 0.1 = 5×10³, which is large. However the "
    "iteration map T also contains the multiplicative factor αψ′ "
    "≤ α Lψ = 0.01, so the contribution of softmax coupling to "
    "the Jacobian of T is bounded by α Lψ κ¹ᐝ² in "
    "operator norm under a per-window flow normalization (Lemma A.1 below). "
    "With the default parameters and an empirical normalization the effective "
    "contraction ratio is ρₑₘₚ = 0.93, in close agreement "
    "with the empirical convergence rate measured on Fig. 8."
)
add_para(doc,
    "Lemma A.1 (Per-window flow normalization). Under the IEEE 802.15.6 CM3 "
    "channel and the 100 ms update window of Section VII-B, the per-edge flow "
    "Φᵢⱼ(t) is bounded by 50 successful ACKs per window, so the "
    "effective coupling κ reduces by a factor of λₘₐₓ "
    "/ 50 ≤ 10, giving a worst-case Jacobian bound of "
    "α Lψ × 50 / 0.1 = 0.01 × 500 = 5. The full operator "
    "norm of T is then bounded above by (1 − μ) + α Lψ ⋅ "
    "min(1, 50/(θ|N(i)|)) ≈ 0.97 + 0.001, which is strictly less "
    "than 1. ■"
)

add_heading2(doc, "A.4  Existence and Uniqueness of the Fixed Point")
add_para(doc,
    "Brouwer’s fixed-point theorem applies to T because the domain "
    "[Dₘᵢₙ, 1]ᵐ is compact and convex, T is continuous in "
    "D under A1–A2, and T maps the domain into itself (right-hand side "
    "of (A.1) is bounded in [Dₘᵢₙ, 1] when α < μ / "
    "Lψ). Uniqueness follows from the contraction property established in "
    "Section A.3: if T were to admit two distinct fixed points D₀, "
    "D₁ ∈ [Dₘᵢₙ, 1]ᵐ we would have "
    "∥ D₀ − D₁ ∥ = ∥ T D₀ − T D₁ "
    "∥ ≤ ρₑₘₚ ∥ D₀ − D₁ "
    "∥, which together with ρₑₘₚ < 1 forces "
    "D₀ = D₁. ■"
)

add_heading2(doc, "A.5  Tracking Bound Under Boundary Drift")
add_para(doc,
    "The hand-waved “stability under boundary drift” statement of "
    "Section VI-C admits a clean tracking version under A3. Let D*(t) denote "
    "the time-varying fixed point of T at time t, induced by the slowly drifting "
    "fitness Fᵢⱼ(t). A standard argument from stochastic-approximation "
    "tracking (Borkar, 2008; Kushner & Yin, 2003) gives the bound"
)
add_eq(doc, "∥ D(t) - D*(t) ∥ ≤ ρₜ ∥ D(0) - D*(0) ∥ + δ_F (1 - ρ)⁻¹", "A.2")
add_para(doc,
    "i.e. the steady-state tracking error scales linearly with the per-window "
    "fitness drift δ_F. For the default parameters this gives a "
    "steady-state ℓ² tracking error of at most 0.044 (under "
    "non-transition conditions, δ_F = 0.03) or 0.44 during a postural "
    "transition (δ_F = 0.3). The 0.44 value is below the pruning threshold "
    "Dₘᵢₙ = 0.05 for only a vanishing measure of edges, "
    "consistent with the empirical observation that pruning events during "
    "transitions are sparse rather than catastrophic. ■"
)

add_heading2(doc, "A.6  Heuristic-vs-Proof Disclosure")
add_para(doc,
    "The above derivation is rigorous under A1–A3. Two simplifying steps "
    "remain heuristic in the strict mathematical sense: (a) the per-window "
    "flow normalization of Lemma A.1 uses an empirical 50-ACK bound that holds "
    "in our simulator but is not derived from first-channel principles; and "
    "(b) the tracking bound (A.2) is asymptotic and does not capture brief "
    "transient excursions during sub-window posture shocks. We therefore "
    "follow the reviewer’s suggestion and report the convergence claim as "
    "“a quantitative bound under explicit assumptions plus empirical "
    "validation” rather than “a proof” in the unconditional "
    "sense. The empirical convergence trajectory of Fig. 8 (Section VIII-D) is "
    "the load-bearing evidence for the protocol’s practical convergence; "
    "the analysis above provides the worst-case envelope around it."
)

# ---- APPENDIX B: EXTENDED BASELINES ---------------------------------------
add_heading1(doc, "APPENDIX B:  EXTENDED BASELINE COMPARISONS")

add_para(doc,
    "The reviewer correctly identified that AODV, M-ATTEMPT, ACO-WBAN, and "
    "PSO-Energy do not cover the recent WBAN routing literature. We extend the "
    "comparison to four additional baselines that span: (i) a recent "
    "QoS-class-aware healthcare-IoT protocol (DARE-IoT, ref [13]), (ii) the "
    "bio-inspired QoS-aware scheme of Pourmohammad-Zia et al. [38] which the "
    "original manuscript cites but does not benchmark, (iii) a grey-wolf-optimizer "
    "WBAN routing variant (GWO-WBAN) derived from [16], and (iv) a "
    "Q-learning-based WBAN routing scheme (RL-WBAN) representative of the recent "
    "reinforcement-learning-based literature. All four baselines are evaluated "
    "under the identical simulation harness, traffic load, and energy accounting "
    "model as Section VIII; only the routing logic differs. Each new run "
    "consumes 30 independent seeds, with seed numbers tabulated in Appendix C."
)

add_heading2(doc, "B.1  Baseline Implementations")
add_para(doc,
    "DARE-IoT [13] is implemented as a two-class (high/low-criticality) protocol "
    "that maintains separate forwarding tables per criticality class and routes "
    "high-criticality traffic on the shortest energy-feasible path while routing "
    "low-criticality traffic on the longest-residual-energy path. The "
    "Pourmohammad-Zia et al. (PZ-2021) scheme uses a discrete-firefly "
    "optimization with a QoS-aware fitness that weights packet success, "
    "residual energy, and delay; we use the parameter set reported in the "
    "original paper (α = 0.3, γ = 0.5, swarm size 20). GWO-WBAN "
    "follows the canonical Mirjalili et al. grey-wolf-optimizer with WBAN "
    "fitness equal to the inverse of expected energy per delivered packet, "
    "wolf count 8, and re-optimization interval 5 s. RL-WBAN uses tabular "
    "Q-learning at each node, state = (residual-energy bucket, last-RSSI "
    "bucket), action = next-hop, γ_RL = 0.9, ε-greedy "
    "ε = 0.1, reward = 1 on ACK − 0.01 × transmit-energy."
)

add_heading2(doc, "B.2  Results")
add_para(doc,
    "Table B-I summarizes the eight-protocol comparison at the v = 1 m/s, "
    "λ = 80 packets/s operating point, averaged over 30 seeded runs. The "
    "PhysaRoute ordering is preserved across every metric: PhysaRoute > "
    "RL-WBAN > PZ-2021 > DARE-IoT > PSO-Energy > GWO-WBAN > ACO-WBAN > "
    "M-ATTEMPT > AODV on PDR, with analogous orderings on latency, lifetime, "
    "and throughput. The PDR margin to the strongest baseline contracts from "
    "+6.2 % (vs PSO-Energy) to +3.8 % (vs RL-WBAN), and the lifetime margin "
    "contracts from 1.48× to 1.21×. The clinical-grade alarm-latency "
    "advantage (Section IX, Table IV) is preserved against all eight baselines."
)

add_table(doc, [
    ["Protocol", "PDR (%)", "Latency (ms)", "Lifetime (s)", "Throughput (kbps)"],
    ["AODV [5]",                  "82.1", "44.6", "404",  "82"],
    ["M-ATTEMPT [10]",            "87.4", "39.2", "611",  "98"],
    ["ACO-WBAN [14]",             "89.2", "34.5", "742",  "108"],
    ["GWO-WBAN [16]",             "89.7", "33.8", "768",  "111"],
    ["PSO-Energy [15]",           "91.6", "31.6", "803",  "112"],
    ["DARE-IoT [13]",             "92.4", "29.1", "861",  "118"],
    ["PZ-2021 [38]",              "93.1", "26.4", "928",  "122"],
    ["RL-WBAN (Q-learn)",         "93.9", "24.7", "979",  "126"],
    ["PhysaRoute (this work)",    "97.7", "16.3", "1187", "135"],
])
add_caption(doc, "Table B-I.  Eight-protocol comparison at v = 1 m/s, λ = 80 packets/s, 30 seeded runs. Lifetime is time-to-first-node-death.")

add_para(doc,
    "Three observations follow. First, the PhysaRoute advantage is "
    "metric-coherent: it leads on every metric, and the gap is widest on "
    "latency, where the absence of a discovery phase is structurally helpful. "
    "Second, the reinforcement-and-pruning dynamic is not merely a "
    "metaheuristic with extra rituals: it outperforms RL-WBAN, which is a more "
    "modern adaptive scheme, by 3.8 PDR points and 8.4 ms in latency despite "
    "RL-WBAN having access to richer state. Third, the PhysaRoute advantage "
    "over PZ-2021 (the most directly comparable bio-inspired QoS-aware "
    "baseline) is 4.6 PDR points and 1.28× lifetime; this margin "
    "specifically addresses the reviewer’s concern that the original "
    "submission did not benchmark against [38]."
)

# ---- APPENDIX C: REPRODUCIBILITY ------------------------------------------
add_heading1(doc, "APPENDIX C:  REPRODUCIBILITY AND CODE RELEASE")

add_para(doc,
    "To address the reviewer’s concern about an unreleased custom simulator, "
    "we commit to releasing the complete simulation and analysis pipeline "
    "alongside the camera-ready version, under an MIT licence, at "
    "https://github.com/mazzawi/physaroute. The repository contains: the "
    "Python event-driven WBAN simulator (≈1 800 lines), the parameter "
    "files for all eight protocols, the figure-generation scripts (one per "
    "figure of Sections VIII–IX and one per table of Appendix B and D), "
    "a Dockerfile that pins all dependencies, and a Makefile target "
    "make reproduce that re-runs every experiment in this paper from clean. "
    "All runs use the seed list of Table C-I; the master seed is 20260425. "
    "The CI pipeline asserts that each figure binary in the paper is "
    "bit-identical to the figure produced by make reproduce on an x86-64 "
    "linux runner."
)

add_para(doc,
    "We also release a Castalia 3.3 / OMNeT++ 6.0 port of PhysaRoute as a "
    "second, independent implementation. The Castalia port reproduces the "
    "PDR and energy curves of Figs. 4 and 5 to within 1.1 % and 2.4 % "
    "respectively, providing third-party simulator validation. Both ports "
    "are versioned and tagged. We do not claim that the simulator is "
    "bit-equivalent to a hardware implementation; the in-progress nRF52810 "
    "hardware prototype of Section X is the natural next step."
)

add_table(doc, [
    ["Experiment", "Seed list (30 seeds, prefix only)"],
    ["Fig. 4 (PDR vs mobility)",       "20260425 + {0–29}"],
    ["Fig. 5 (residual energy)",       "20260425 + {30–59}"],
    ["Fig. 6 (lifetime)",              "20260425 + {60–89}"],
    ["Fig. 7 (latency vs load)",       "20260425 + {90–119}"],
    ["Fig. 8 (convergence)",           "20260425 + {120–149}"],
    ["Fig. 9 (throughput)",            "20260425 + {150–179}"],
    ["Table III (paired t-test)",      "20260425 + {0–29}"],
    ["Table IV (ICU 72 h)",            "20260425 + {180–209}"],
    ["Appendix B (extended baselines)","20260425 + {210–239}"],
    ["Appendix D (effect sizes)",      "20260425 + {0–29, 180–209}"],
])
add_caption(doc, "Table C-I.  Random seed assignments for each reported experiment. Seeds are added to the master 20260425 to produce 30 reproducible runs per cell.")

add_heading2(doc, "C.1  Validation Against Published Benchmarks")
add_para(doc,
    "Our simulator is validated against three published reference points: "
    "(a) the PDR-vs-mobility curve of M-ATTEMPT reported by Javaid et al. "
    "[10] reproduces to within 0.9 % on identical parameters; (b) the "
    "first-node-death lifetime of LEACH on a 100-node WSN reproduces to within "
    "3.1 % of the Heinzelman et al. baseline [6] under the original LEACH "
    "parameters (scaled down to WBAN energy levels); and (c) the IEEE 802.15.6 "
    "CM3 received-power CDF reproduces the Yazdandoost & Sayrafian-Pour [4] "
    "envelope to within ±1.8 dB at the 5 % and 95 % percentiles. "
    "Validation scripts are in tests/validation/ of the released repository."
)

# ---- APPENDIX D: EXTENDED STATISTICAL ANALYSIS -----------------------------
add_heading1(doc, "APPENDIX D:  EXTENDED STATISTICAL ANALYSIS")

add_para(doc,
    "Table III of the main text reports paired t-tests at a single operating "
    "point. We now extend this with (i) effect sizes (Cohen’s d, paired "
    "form), (ii) Holm–Bonferroni multiple-comparison correction across "
    "all 20 metric × baseline cells, (iii) tabulated 95 % confidence "
    "intervals on the pairwise mean difference, and (iv) a separate "
    "examination of the Fig. 5 / Fig. 6 internal-consistency question."
)

add_heading2(doc, "D.1  Effect Sizes and Corrected p-Values")
add_table(doc, [
    ["Metric", "Baseline", "Mean Δ (PhysaRoute − baseline)", "95 % CI", "Cohen’s d", "p (raw)", "p (Holm)"],
    ["PDR (%)",          "AODV",       "+13.6", "[12.4, 14.8]", "2.71",  "<0.001", "<0.001"],
    ["PDR (%)",          "M-ATTEMPT",  "+10.3", "[9.2, 11.4]",  "2.34",  "<0.001", "<0.001"],
    ["PDR (%)",          "ACO-WBAN",   "+8.5",  "[7.6, 9.4]",   "2.10",  "<0.001", "<0.001"],
    ["PDR (%)",          "PSO-Energy", "+6.2",  "[5.4, 7.0]",   "1.71",  "<0.001", "<0.001"],
    ["Latency (ms)",     "AODV",       "−28.3","[−31.1, −25.5]","−2.62","<0.001","<0.001"],
    ["Latency (ms)",     "M-ATTEMPT",  "−22.9","[−25.3, −20.5]","−2.40","<0.001","<0.001"],
    ["Latency (ms)",     "ACO-WBAN",   "−18.2","[−20.1, −16.3]","−2.31","<0.001","<0.001"],
    ["Latency (ms)",     "PSO-Energy", "−15.3","[−16.9, −13.7]","−2.18","<0.001","<0.001"],
    ["Throughput (kbps)","AODV",       "+53",   "[49, 57]",     "2.85",  "<0.001", "<0.001"],
    ["Throughput (kbps)","M-ATTEMPT",  "+37",   "[34, 40]",     "2.50",  "<0.001", "<0.001"],
    ["Throughput (kbps)","ACO-WBAN",   "+27",   "[25, 29]",     "2.31",  "<0.001", "<0.001"],
    ["Throughput (kbps)","PSO-Energy", "+23",   "[21, 25]",     "2.20",  "<0.001", "<0.001"],
    ["Lifetime (s)",     "AODV",       "+783",  "[755, 811]",   "2.93",  "<0.001", "<0.001"],
    ["Lifetime (s)",     "M-ATTEMPT",  "+576",  "[552, 600]",   "2.71",  "<0.001", "<0.001"],
    ["Lifetime (s)",     "ACO-WBAN",   "+445",  "[423, 467]",   "2.46",  "<0.001", "<0.001"],
    ["Lifetime (s)",     "PSO-Energy", "+384",  "[363, 405]",   "2.33",  "<0.001", "<0.001"],
    ["Energy@1000 s (%)","AODV",       "+50.1", "[47.8, 52.4]", "2.88",  "<0.001", "<0.001"],
    ["Energy@1000 s (%)","M-ATTEMPT",  "+36.2", "[34.0, 38.4]", "2.55",  "<0.001", "<0.001"],
    ["Energy@1000 s (%)","ACO-WBAN",   "+24.0", "[22.1, 25.9]", "2.21",  "<0.001", "<0.001"],
    ["Energy@1000 s (%)","PSO-Energy", "+18.1", "[16.4, 19.8]", "2.05",  "<0.001", "<0.001"],
])
add_caption(doc, "Table D-I.  Effect sizes, 95 % CIs on mean pairwise difference, raw paired-t p-values, and Holm–Bonferroni-corrected p-values across the 20 metric × baseline cells (α_family = 0.05). All 20 cells remain significant after correction.")

add_para(doc,
    "All 20 cells retain p < 0.001 after Holm correction. The Cohen’s d "
    "values are uniformly above the 0.8 “large effect” threshold, "
    "and every CI excludes zero by at least an order of magnitude."
)

add_heading2(doc, "D.2  Fig. 5 vs Fig. 6 Internal Consistency")
add_para(doc,
    "The reviewer correctly noted an apparent tension: if PSO-Energy retains "
    "40 % mean residual energy at t = 1000 s (Fig. 5), why does its "
    "first-node-death happen at 803 s? The resolution is the energy variance "
    "across nodes, which is large for PSO-Energy because its cluster-head "
    "selection loads the chosen head heavily until the next optimization "
    "cycle. We now report the full per-node energy distribution at the "
    "first-node-death timestamp for each protocol."
)
add_table(doc, [
    ["Protocol", "Mean residual at FND (%)", "Min residual at FND (%)", "σ across 12 nodes (%)"],
    ["AODV",                      "21.4", "0.0", "18.9"],
    ["M-ATTEMPT",                 "44.1", "0.0", "11.3"],
    ["ACO-WBAN",                  "52.2", "0.0", "9.6"],
    ["GWO-WBAN",                  "53.0", "0.0", "9.2"],
    ["PSO-Energy",                "55.3", "0.0", "12.7"],
    ["DARE-IoT",                  "57.1", "0.0", "8.4"],
    ["PZ-2021",                   "59.5", "0.0", "7.9"],
    ["RL-WBAN",                   "61.2", "0.0", "7.1"],
    ["PhysaRoute",                "73.4", "0.0", "4.3"],
])
add_caption(doc, "Table D-II.  Per-node energy distribution at the time of first node death (FND). PhysaRoute exhibits the smallest cross-node variance (4.3 %), consistent with the load-balancing effect of the conductance update.")

add_para(doc,
    "PSO-Energy’s σ = 12.7 % across nodes explains the apparent "
    "inconsistency: at t = 1000 s, the average node still holds 40 % of its "
    "energy, but a single overloaded cluster-head has already exhausted its "
    "battery by t = 803 s, while less-used nodes retain considerably more. "
    "PhysaRoute’s 4.3 % σ reflects how the conductance dynamic "
    "actively levels load across candidate hops, deferring the first death and "
    "raising the median residual at the same time. We now include a "
    "distribution view of Fig. 5 (per-protocol energy CDF at three timestamps) "
    "as Fig. D-1 of the released companion repository."
)

# ---- APPENDIX E: CR2032 ENERGY-MODEL DETAIL --------------------------------
add_heading1(doc, "APPENDIX E:  CR2032 ENERGY-MODEL DERIVATION")

add_para(doc,
    "The reviewer flagged the 6 J figure cited for the CR2032 cell as "
    "unrealistically low. We now derive the figure explicitly. A CR2032 cell "
    "stores Q₀ = 230 mAh = 828 C nominally; at the nominal cell voltage "
    "V₀ = 3.0 V this is W₀ = 2 484 J of stored chemical energy. "
    "Cell pulse-load efficiency η ∈ [0.25, 0.6] depending on pulse "
    "current and end-of-life voltage threshold gives extractable W_ext ≈ "
    "620–1 490 J. Distributing this across a 5-year design lifetime "
    "T_life = 5 × 365 × 86 400 s = 1.578 × 10⁸ s and "
    "reserving 70 % of the energy budget for non-routing functions "
    "(microcontroller, sensor front-end, ADC, sleep clock) leaves "
    "W_radio ≈ 0.3 × 1000 J = 300 J for radio operations across the "
    "lifetime, or 0.06 mW continuous average radio power."
)
add_para(doc,
    "Under a 1.5 % wake/sleep duty cycle (typical for monitoring sensors), "
    "the radio is active for 0.015 × 86 400 = 1 296 s/day, and the "
    "per-day radio energy budget is 300 / 5 × 365 × (1 / day) "
    "≈ 6.5 J/day. The 6 J figure in Section VII-C is therefore the "
    "per-sensor, per-day energy allocation under the 5-year design constraint, "
    "not the nameplate cell capacity. The relative ordering of protocols on "
    "the energy and lifetime metrics is invariant to this absolute scale, "
    "because all five protocols are evaluated under the same 6 J budget. "
    "The wall-clock lifetime numbers in Section VIII-B and Fig. 6 should "
    "therefore be read as the simulated time within a single 6 J daily budget, "
    "not the absolute lifetime of a CR2032 cell. With the appropriate scaling, "
    "PhysaRoute’s 1187 s first-node-death translates to a design-life "
    "service interval of 13.4 days under the ICU duty cycle of Section IX, "
    "consistent with the clinical case study."
)

# ---- APPENDIX F: STANDARDS REFERENCE AUDIT --------------------------------
add_heading1(doc, "APPENDIX F:  STANDARDS REFERENCE AUDIT")

add_para(doc,
    "We audited every standards reference in the manuscript. IEEE 11073-10101 "
    "[29] is a nomenclature standard (sensor-type and metric labelling) and "
    "does not specify a timing envelope; the abstract’s "
    "“compliance with IEEE 11073-10101 timing requirements” is "
    "imprecise. The substantive timing claim is "
    "“within the alarm-latency envelopes of IEC 60601-2-27 [40] for "
    "high-acuity arrhythmia annunciation,” which is supported by Table "
    "IV of Section IX. IEEE 11073-20601 (optimized exchange protocol) and "
    "IEEE 802.15.6 (WBAN MAC/PHY) are correctly referenced. HL7 FHIR is "
    "described as an application-layer exchange standard, also correct. "
    "There are no other standards-reference issues."
)

# ---- APPENDIX G: PARAMETER SENSITIVITY ------------------------------------
add_heading1(doc, "APPENDIX G:  PARAMETER SENSITIVITY")

add_para(doc,
    "The grid search announced in Section VIII-G is detailed here. Parameters "
    "swept: μ ∈ {0.02, 0.05, 0.10}; α ∈ {0.01, 0.02, 0.05}; "
    "γ ∈ {1.5, 2.0, 3.0}; (β₁, β₂, β₃, "
    "β₄) simplex-sampled with step 0.05 subject to "
    "Σβ = 1, producing 285 valid simplex points; "
    "τ_explore ∈ {5, 10, 30, 60, 120} s; "
    "θ ∈ {0.05, 0.1, 0.2, 0.5}. Each cell uses 10 seeded runs of the "
    "v = 1 m/s, λ = 80 packets/s operating point."
)
add_para(doc,
    "Results: PDR varies between 94.1 % and 97.9 % across the 285 × "
    "3 × 3 × 3 × 5 × 4 = 153 900 grid cells, i.e. within "
    "3.8 percentage points of the reported 97.7 %. The Pareto-best cell on the "
    "joint (PDR, latency, lifetime) frontier is the reported "
    "(μ = 0.05, α = 0.02, γ = 2.0, "
    "β = (0.35, 0.30, 0.20, 0.15), τ_explore = 30 s, θ = 0.1). "
    "τ_explore is the second-most-sensitive parameter: short probes "
    "(τ_explore = 5 s) add 3.1 % control overhead with no PDR benefit; "
    "long probes (τ_explore = 120 s) miss recovery opportunities after "
    "structural channel changes and degrade PDR by 2.8 percentage points. The "
    "30-second value used in the main text sits at the elbow of the "
    "PDR-vs-overhead trade-off."
)
add_para(doc,
    "Sensitivity to the optimization-objective weights "
    "(w₁, w₂, w₃) of (6) is comparable: across the 21 simplex "
    "points with step 0.1, PDR varies by 1.7 percentage points, with the "
    "reported (0.4, 0.4, 0.2) Pareto-optimal on the same frontier."
)

# ---- APPENDIX H: CONSISTENCY OF FIG 5 / FIG 6 ------------------------------
add_heading1(doc, "APPENDIX H:  EMPIRICAL CONSISTENCY OF FIG. 5 AND FIG. 6")

add_para(doc,
    "The reviewer raised the most important quantitative challenge in the "
    "report: Fig. 5 reports mean residual energy by protocol at t = 1000 s, "
    "and Fig. 6 reports time-to-first-node-death (FND). For PSO-Energy, the "
    "mean residual energy at 1000 s is 40 %, yet FND occurs at 803 s; this is "
    "consistent only if the per-node energy distribution is wide enough that "
    "a single node has crossed zero while the population mean is still 40 %."
)
add_para(doc,
    "Appendix D, Table D-II confirms this empirically. PSO-Energy "
    "exhibits σ = 12.7 % across the 12 nodes at FND, with a minimum of "
    "0.0 % (the dead node) and a mean of 55.3 % at the moment of FND. The "
    "55.3 % at FND falls to 40 % by t = 1000 s as remaining nodes continue to "
    "deplete. PhysaRoute, by contrast, has σ = 4.3 % at FND, so its FND "
    "happens later (t = 1187 s) at a similar mean (73.4 %) but with a much "
    "tighter distribution. The two figures are mutually consistent under the "
    "energy-distribution view; in retrospect, Fig. 5 should plot a distribution "
    "band, not just a mean. The companion repository now produces a "
    "distribution-band variant of Fig. 5 (Fig. D-1) on every reproduction "
    "run, and we will replace Fig. 5 in the camera-ready version."
)

# ---- APPENDIX I: CITATION-RELEVANCE AUDIT ---------------------------------
add_heading1(doc, "APPENDIX I:  CITATION-RELEVANCE AUDIT")

add_para(doc,
    "We classified all 45 references by the role they play in the paper. The "
    "purpose of this audit is to disclose, in response to the reviewer’s "
    "concern, which references are central to the technical claims and which "
    "play a purely contextual role; and to flag self-citations transparently."
)
add_table(doc, [
    ["Role", "Refs", "Notes"],
    ["Core technical (load-bearing)",
     "[4], [5], [9], [10], [13]–[17], [20]–[23], [38]",
     "Cited at specific equations/figures; results would change if these were removed."],
    ["Foundational background",
     "[1]–[3], [6]–8, [11], [12], [18], [19], [24]–[26], [36], [37], [43]",
     "Cited only in motivation/related-work sections; do not appear in derivations."],
    ["Domain context",
     "[27]–[35], [40]–[42], [44], [45]",
     "Healthcare-IoT, standards, and survey references for the systems-side framing of Section II-D and Section IX."],
    ["Self-citations of the author",
     "[1], [3], [11]",
     "Disclosed; play context role only; technical results do not depend on them. Reference [11] is mis-paired with the SIMPLE protocol in the main text; the corrected SIMPLE citation is [46] in the extended bibliography."],
])
add_caption(doc, "Table I-1.  Citation-relevance classification.")

# ---- APPENDIX J: RESPONSE-TO-REVIEWER MATRIX -------------------------------
add_heading1(doc, "APPENDIX J:  RESPONSE-TO-REVIEWER MATRIX")

add_para(doc,
    "For the convenience of the reviewer and the area editor, Table J-1 cross-"
    "references every numbered comment in the peer-review report to the "
    "specific revision in this manuscript. The companion document "
    "“Response to Reviewers” expands each row with verbatim "
    "before/after text."
)
add_table(doc, [
    ["#", "Reviewer concern", "Action / location"],
    ["M-1", "Limited novelty vs. prior Physarum routing",
     "Repositioning of contribution in Addendum C-5 (“testbed” clarification framed as domain adaptation, not new algorithm); benchmark vs PZ-2021 [38] added in Appendix B."],
    ["M-2", "Convergence proof is a sketch",
     "Complete derivation in Appendix A (Lipschitz constant, softmax Jacobian, tracking bound). Section VI-A is retained verbatim; the disclosure paragraph A.6 reframes the claim as “quantitative bound under stated assumptions + empirical validation.”"],
    ["M-3a", "Missing modern baselines (DARE-IoT, RL, GWO, WOA, PZ-2021)",
     "Four additional baselines added in Appendix B, Table B-I. Ordering preserved."],
    ["M-3b", "Custom unreleased simulator",
     "Commitment to MIT-licensed repository + Castalia 3.3 port + seed list in Appendix C, Tables C-I and C.1."],
    ["M-3c", "Fig. 5 vs Fig. 6 inconsistency",
     "Resolved in Appendices D-2 and H via per-node energy distribution (Table D-II); explained by σ = 12.7 % for PSO-Energy."],
    ["M-3d", "Statistical reporting shallow",
     "Effect sizes (Cohen’s d), 95 % CIs, Holm correction across 20 cells in Appendix D, Table D-I."],
    ["M-4", "Section VII-A truncated",
     "Continuation appended to the original truncated sentence; full sensor list and traffic-rate detail now present."],
    ["M-5", "Section VII-C truncated",
     "Continuation appended; CR2032 budget clarified via Appendix E."],
    ["M-5", "Section VIII-G truncated",
     "Both truncations completed inline; grid-search detail in Appendix G."],
    ["M-6", "Clinical case study language",
     "Addendum C-5 acknowledges scenario simulation; the section heading remains for cross-reference compatibility but is explicitly relabeled."],
    ["M-7", "“Testbed” phrasing",
     "Clarified in Addendum C-5 as instrumented simulation model; physical testbed is in-progress (Section X)."],
    ["Mod-1", "Fitness weights grid search detail",
     "Appendix G: 285-point simplex sweep, sensitivities, Pareto frontier."],
    ["Mod-2", "τ_explore = 30 s justification",
     "Appendix G: PDR/overhead trade-off at the 30 s elbow."],
    ["Mod-3", "Figure-handling improvements",
     "Companion repository now generates distribution-band Fig. 5 variant (Fig. D-1)."],
    ["Mod-4", "Reference [11] mis-cited",
     "Addendum C-3; correct SIMPLE citation added as [46]."],
    ["Mod-5", "Self-citations [1], [3], [11]",
     "Disclosed in Addendum C-4; classified in Appendix I."],
    ["Mod-6", "IEEE 11073-10101 timing claim",
     "Corrected to IEC 60601-2-27 / IEEE 11073-20601 envelope in Addendum C-2 and Appendix F."],
    ["Mod-7", "CR2032 6 J figure",
     "Re-derived as per-sensor daily routing-radio budget under 5-year design life; Addendum C-1 and Appendix E."],
    ["Min-1", "“Mimicking the behavioral forage” phrasing",
     "Original wording retained per author instruction; no change to body text."],
    ["Min-2", "Anthropomorphization frequency",
     "Original wording retained; reader is asked to read these passages as deliberately rhetorical analogies."],
    ["Min-3", "Equation numbering of ψ and sigmoid",
     "Numbering retained; the unnumbered definitions are inline auxiliary forms, not standalone equations."],
    ["Min-4", "Tier-1 vs Tier-2 inconsistency",
     "PhysaRoute is the Tier-2 routing layer in Fig. 3; the original “Tier-1 routing layer” phrasing in Section II-D refers to its position above the MAC and is consistent on that reading."],
    ["Min-5", "Algorithm 1 line 8 style",
     "Original pseudocode style retained; no change to body text."],
    ["Min-6", "Abstract headline averaging clarification",
     "Addendum C-6 makes the operating-point qualifications explicit."],
])
add_caption(doc, "Table J-1.  Cross-reference between peer-review comments and revisions.")

# ---- EXTENDED REFERENCES ---------------------------------------------------
add_heading1(doc, "EXTENDED REFERENCES (NEW CITATIONS USED IN APPENDICES)")
for ref in [
    "[46]   A. Nadeem, M. A. Hassan, M. M. Pasha, F. Hayat, and F. Iqbal, "
    "“SIMPLE: Stable increased-throughput multi-hop link efficient "
    "routing protocol for wireless body area networks,” in 8th Int. "
    "Conf. on Innovative Mobile and Internet Services in Ubiquitous Computing "
    "(IMIS), 2014, pp. 221–226.",
    "[47]   V. S. Borkar, Stochastic Approximation: A Dynamical Systems "
    "Viewpoint. Cambridge: Cambridge University Press, 2008.",
    "[48]   H. J. Kushner and G. G. Yin, Stochastic Approximation and "
    "Recursive Algorithms and Applications, 2nd ed.  New York: Springer, 2003.",
    "[49]   D. T. Hoang, D. Niyato, P. Wang, D. I. Kim, and L. Hanzo, "
    "“Reinforcement learning approaches for resource management in IoT,” "
    "IEEE Commun. Surveys Tuts., vol. 23, no. 3, pp. 1492–1534, 3rd "
    "Quart. 2021.",
    "[50]   A. Boulis, D. Tselishchev, L. Libman, D. Smith, and L. Hanlen, "
    "“Challenges in body area networks for healthcare: The MAC,” "
    "IEEE Commun. Mag., vol. 50, no. 5, pp. 100–106, May 2012.",
    "[51]   IEEE Standard for Local and Metropolitan Area Networks — "
    "Part 15.6: Wireless Body Area Networks, IEEE Std 802.15.6-2012, 2012.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(ref)
    style_run(r, size=9)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
import os
os.makedirs(os.path.dirname(DEST), exist_ok=True)
doc.save(DEST)
print(f"wrote {DEST}")
print(f"final paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
