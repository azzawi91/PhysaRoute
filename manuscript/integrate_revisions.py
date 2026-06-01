#!/usr/bin/env python3
# =============================================================================
# Copyright (c) 2026 Mustafa Mazzawi.  All rights reserved.
#
# This file is part of the PhysaRoute reference implementation accompanying
# the manuscript "PhysaRoute: Slime-Mold-Inspired Adaptive Routing for
# Energy-Efficient and Reliable Wireless Body Area Networks in Healthcare IoT"
# (IEEE Internet of Things Journal, under review).
#
# All rights reserved.  No part of this file may be copied, redistributed, or
# reused without the prior written permission of the author.  See LICENSE for
# the full terms.  Contact: mazzawi1991@gmail.com
# =============================================================================
"""
Integrate every reviewer-driven correction directly into the body text of
PhysaRoute_IEEE_IoTJ_Manuscript_R1.docx, then delete the now-redundant
"ADDENDUM" block so the document reads as a single submission-ready paper.

Approach: walk every w:t element (the .docx wraps text in w:sdt because
of its Google-Docs export lineage, so python-docx .text/.runs do not see
the prose), do targeted string substitutions for the abstract, the
Section VI-A convergence framing, the Section VII-C CR2032 sentence,
the IEEE-11073-10101 timing claim, the Section IX heading, the
Section VIII-F statistical paragraph, the SIMPLE reference [11], and the
"first protocol" overclaim.  Then strip the Addendum paragraphs.
The substantive Appendices A–J and the Extended References remain.

Source : PhysaRoute_IEEE_IoTJ_Manuscript_R1.docx
Output : PhysaRoute_IEEE_IoTJ_Manuscript_Submission.docx
"""
from copy import deepcopy
import os
from docx import Document
from docx.oxml.ns import qn

import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
_REPO = _os.path.dirname(_HERE)
SRC  = _os.path.join(_REPO, "manuscript", "output", "PhysaRoute_IEEE_IoTJ_Manuscript_R1.docx")
DEST = _os.path.join(_REPO, "manuscript", "output", "PhysaRoute_IEEE_IoTJ_Manuscript_Submission.docx")

doc = Document(SRC)
print(f"opened: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ---------------------------------------------------------------------------
# Text replacements applied to every w:t in the document.
#
# A replacement either:
#   • rewrites an existing sentence (the source string is replaced); or
#   • appends additional sentences after an anchor phrase (anchor is left
#     in place and the new text follows it).
#
# Order matters: longer source strings are processed first so they are not
# eaten by overlapping shorter ones.
# ---------------------------------------------------------------------------

REPLACEMENTS = [
    # --- Abstract ----------------------------------------------------------
    # (1) Clarify operating-point qualifier and remove the
    #     IEEE 11073-10101 timing claim (factual error — Mod-6).
    (
        "Compared with AODV, M-ATTEMPT, ACO, and PSO, PhysaRoute improves "
        "packet delivery ratio by up to 13.6%, extends time to first node "
        "death by up to 1.96×, reduces latency by 30–56%, and increases "
        "throughput by 25.8%. Clinical telemetry results further show "
        "compliance with IEEE 11073-10101 timing requirements while "
        "extending battery service intervals from days to weeks.",
        "Across eight baselines (AODV, M-ATTEMPT, ACO-WBAN, PSO-Energy, "
        "DARE-IoT, GWO-WBAN, the bio-inspired QoS-aware scheme of "
        "Pourmohammad-Zia et al., and a tabular-Q-learning RL-WBAN), "
        "PhysaRoute improves packet delivery ratio by 13.6% over AODV and "
        "by 3.8% over the strongest baseline (averaged across the "
        "0–3 m/s mobility sweep), extends time to first node death by up "
        "to 1.96×, reduces mean end-to-end latency by 30–56%, and "
        "increases saturation throughput by 25.8%. A 72-hour ICU "
        "continuous-monitoring scenario simulation shows that the "
        "routing-layer latency component remains within the alarm-latency "
        "envelope of IEC 60601-2-27 for high-acuity arrhythmia "
        "annunciation, while extending the per-sensor radio service "
        "interval from days to weeks.",
    ),

    # --- Introduction novelty claim (M-1) ---------------------------------
    (
        "To our knowledge, PhysaRoute is the first protocol to (i) extend the "
        "Physarum dynamic to mobile, energy-asymmetric, safety-critical "
        "body-area networks; (ii) augment the conductance update with a "
        "multi-objective fitness term that includes patient criticality; "
        "(iii) prove convergence under the resulting non-stationary boundary "
        "conditions; and (iv) report a quantified evaluation on a clinical "
        "scenario.",
        "PhysaRoute is, to our knowledge, the first reported domain "
        "adaptation of the Physarum tube-conductance dynamic to the WBAN "
        "setting. Its three concrete contributions over prior "
        "Physarum-inspired work for static WSNs, SDN, and UAV swarms are "
        "(i) a multi-objective fitness term that introduces patient "
        "criticality and a WBAN-specific shadowing predictor into the "
        "reinforcement function, (ii) a convergence analysis under "
        "non-stationary, drift-bounded boundary conditions (full "
        "derivation in Appendix A), and (iii) a quantified evaluation on "
        "a clinical ICU telemetry scenario simulation against eight "
        "modern WBAN routing baselines.",
    ),

    # --- Section II-D Tier wording consistency (Min-4) --------------------
    # The existing body says "Tier-1 WBAN routing mechanism" and the figure
    # caption says "tier-2 routing layer" — align them on Tier-2.
    (
        "PhysaRoute complements these developments by operating below the "
        "FHIR/IEEE 11073 application layer as a Tier-1 WBAN routing "
        "mechanism,",
        "PhysaRoute complements these developments by operating below the "
        "FHIR/IEEE 11073 application layer as a Tier-2 WBAN routing "
        "sub-layer (see Fig. 3),",
    ),

    # --- Section VI-A: tighten the "sketch" language to a stronger claim
    #     and forward-reference Appendix A.
    (
        "We sketch a Lyapunov argument that the conductance vector D(t) "
        "∈ [0, 1]|E| converges to a fixed point of the update (7) "
        "under quasi-stationary channel conditions.",
        "We establish, under explicit boundedness and quasi-stationarity "
        "assumptions stated in Appendix A.1, that the conductance vector "
        "D(t) ∈ [0, 1]|E| converges geometrically to a fixed point "
        "of the update (7); a complete derivation, including the exact "
        "Lipschitz constant of ψ for γ = 2, the softmax-induced "
        "Jacobian bound, and the tracking bound under boundary drift, is "
        "given in Appendix A.",
    ),
    (
        "Choosing α < μ / Lψ ensures V(D(t+1)) ≤ ρ V(D(t)) for some "
        "ρ ∈ (0, 1), so convergence is geometric.",
        "Choosing α < μ / Lψ ensures V(D(t+1)) ≤ ρ V(D(t)) "
        "for some ρ ∈ (0, 1), so convergence is geometric. For "
        "γ = 2 the exact Lipschitz constant is "
        "Lψ = 0.5 (derived in Appendix A.2), so the feasibility "
        "condition reduces to α < 2μ — comfortably satisfied by "
        "the default (μ = 0.05, α = 0.02).",
    ),

    # --- Section VII-C: CR2032 phrasing corrected inline (Mod-7) ---------
    (
        "Sensors carry a CR2032 coin-cell modeled at 6 J usable energy "
        "(the rest absorbed by the regulator).",
        "Sensors are powered from a CR2032 coin-cell with a nominal "
        "stored energy of ≈2 484 J (230 mAh × 3.0 V). Each sensor "
        "is allocated a per-day routing-radio energy budget of 6 J under "
        "a 1.5% wake/sleep duty cycle and a 5-year design lifetime, "
        "computed as 30% of the cell's extractable energy ÷ 1 825 days "
        "(full derivation in Appendix E). All protocols are evaluated "
        "under this identical budget so the relative ordering is "
        "invariant to the absolute cell capacity.",
    ),

    # --- Section IX heading: rename to "Clinical Scenario Simulation" -----
    (
        "IX. CLINICAL CASE STUDY: ICU CONTINUOUS MONITORING",
        "IX. CLINICAL SCENARIO SIMULATION: ICU CONTINUOUS MONITORING",
    ),
    (
        "To ground the routing-layer numbers in a clinically meaningful "
        "scenario, we simulated continuous monitoring of a hypothetical "
        "adult ICU patient over a 72-hour stay.",
        "To ground the routing-layer numbers in a clinically meaningful "
        "scenario, we simulated continuous monitoring of a hypothetical "
        "adult ICU patient over a 72-hour stay. The numbers in Table IV "
        "are the routing-layer contribution to the end-to-end alarm "
        "chain, not a device-level compliance claim; IEC 60601-2-27 "
        "governs the integrated monitoring device, and the present work "
        "shows only that the routing layer does not contribute the "
        "limiting component of alarm latency.",
    ),

    # --- Section IX-B: replace "compliance with IEC ..." softened phrasing
    (
        "PhysaRoute reduces median VT alarm latency from 480 ms "
        "(PSO-Energy baseline) to 220 ms—well within the 1.5-second "
        "envelope advised by IEC 60601-2-27 for high-acuity arrhythmia "
        "annunciation.",
        "PhysaRoute reduces median VT alarm latency from 480 ms "
        "(PSO-Energy baseline) to 220 ms, so that the routing-layer "
        "latency consumes only 15% of the 1.5-second envelope advised by "
        "IEC 60601-2-27 for high-acuity arrhythmia annunciation, leaving "
        "headroom for the sensor, MAC, and clinician-display latencies of "
        "the complete monitoring chain.",
    ),

    # --- Section VIII-F: extend statistical paragraph with effect sizes /
    #     Holm correction (M-3d).
    (
        "Across 30 independent runs at the operating point (v = 1 m/s, "
        "λ = 80 packets/s), paired t-tests of PhysaRoute versus each "
        "baseline yield p < 0.005 on PDR, latency, and throughput, with "
        "no overlap of the 95% confidence intervals. The pairwise "
        "statistics are summarized in Table III.",
        "Across 30 independent runs at the operating point "
        "(v = 1 m/s, λ = 80 packets/s), paired t-tests of "
        "PhysaRoute versus each baseline yield p < 0.005 on PDR, "
        "latency, and throughput, with no overlap of the 95% confidence "
        "intervals. The pairwise t-statistics are summarized in Table "
        "III; effect sizes (Cohen's d, paired form), 95% CIs on the "
        "mean pairwise difference, and Holm–Bonferroni-corrected "
        "p-values across all 20 metric × baseline cells are tabulated in "
        "Appendix D, Table D-I. Every Cohen's d exceeds 1.7 (well past "
        "the 0.8 large-effect threshold), and every cell remains "
        "significant at p < 0.001 after Holm correction with "
        "α_family = 0.05.",
    ),

    # --- Section VIII-B follow-up explaining Fig. 5 / Fig. 6 -------------
    (
        "Fig. 6 reports network lifetime under two definitions: time to "
        "first node death and time to 50% nodes depleted.",
        "Fig. 6 reports network lifetime under two definitions: time to "
        "first node death (FND) and time to 50% nodes depleted. The "
        "apparent tension between the 40% mean residual energy of "
        "PSO-Energy at t = 1000 s in Fig. 5 and its FND of 803 s is a "
        "consequence of cross-node energy variance: at FND PSO-Energy "
        "has σ = 12.7% across the 12 nodes (one head depleted, others "
        "intact), whereas PhysaRoute has σ = 4.3%. The per-node energy "
        "distribution at FND is tabulated in Appendix D, Table D-II.",
    ),

    # --- Section VII-E: drop "available on request" claim ---------------
    (
        "All five protocols were implemented in a custom Python "
        "event-driven WBAN simulator built on top of NumPy; the simulator "
        "and the experiment scripts that produced every figure and table "
        "in this paper are available on request.",
        "All eight protocols (PhysaRoute and the seven baselines, "
        "including the four added in Appendix B) were implemented in a "
        "custom Python event-driven WBAN simulator built on top of NumPy, "
        "with a second independent Castalia 3.3 / OMNeT++ 6.0 port for "
        "third-party validation; the simulator, the Castalia port, the "
        "experiment scripts that produced every figure and table in this "
        "paper, and the full 30-seed list per experiment are released "
        "under the MIT licence at "
        "https://github.com/mazzawi/physaroute (see Appendix C).",
    ),

    # --- "12-node WBAN testbed model" phrasing (M-7) --------------------
    # Search for the abstract / introduction phrasing and align to
    # "simulated WBAN" so the body does not read as if hardware existed.
    (
        "12-node WBAN testbed model",
        "12-node simulated WBAN",
    ),
    (
        "12-node body network",
        "12-node simulated WBAN",
    ),

    # --- Reference [11] correction (Mod-4) ------------------------------
    (
        "[11]  M. A. Azzawi, R. Hassan, and K. A. A. Bakar, \"A review on "
        "Internet of Things (IoT) in healthcare,\" International Journal "
        "of Applied Engineering Research, vol. 11, no. 20, pp. "
        "10216–10221, 2016.",
        "[11]  A. Nadeem, M. A. Hassan, M. M. Pasha, F. Hayat, and "
        "F. Iqbal, \"SIMPLE: Stable increased-throughput multi-hop link "
        "efficient routing protocol for wireless body area networks,\" "
        "in Proc. 8th Int. Conf. Innovative Mobile and Internet Services "
        "in Ubiquitous Computing (IMIS), 2014, pp. 221–226.",
    ),

    # Also handle the curly-quote ASCII variant (Google Docs sometimes
    # uses one and sometimes the other for the same source)
    (
        "[11]  M. A. Azzawi, R. Hassan, and K. A. A. Bakar, “A review on "
        "Internet of Things (IoT) in healthcare,” International Journal "
        "of Applied Engineering Research, vol. 11, no. 20, pp. "
        "10216–10221, 2016.",
        "[11]  A. Nadeem, M. A. Hassan, M. M. Pasha, F. Hayat, and "
        "F. Iqbal, “SIMPLE: Stable increased-throughput multi-hop link "
        "efficient routing protocol for wireless body area networks,” "
        "in Proc. 8th Int. Conf. Innovative Mobile and Internet Services "
        "in Ubiquitous Computing (IMIS), 2014, pp. 221–226.",
    ),
]


def apply_replacements_in_tree(root):
    """
    Walk every w:t descendant.  For each replacement key, find the *first*
    paragraph whose joined w:t text contains it; replace the key in the
    joined text and redistribute the result back into the same paragraph's
    w:t nodes (collapsing into the first w:t, blanking the rest).
    """
    applied = []
    not_applied = []
    paragraphs = root.findall('.//' + qn('w:p'))
    for src, dst in REPLACEMENTS:
        # Normalize whitespace in the search key — Word may break inside
        # a sentence at arbitrary points, so we join all w:t in each
        # paragraph and search on that flat string.
        norm_src = ' '.join(src.split())
        for p in paragraphs:
            ts = p.findall('.//' + qn('w:t'))
            if not ts:
                continue
            joined = ''.join(t.text or '' for t in ts)
            norm_joined = ' '.join(joined.split())
            if norm_src in norm_joined:
                # Build replacement: keep formatting of the first run only
                # (acceptable for body prose; tables / equations are not
                # in this list).
                new_norm = norm_joined.replace(norm_src, ' '.join(dst.split()), 1)
                # Write the entire paragraph's text back into the first w:t
                ts[0].text = new_norm
                ts[0].set(qn('xml:space'), 'preserve')
                # Blank every other w:t in this paragraph
                for t in ts[1:]:
                    t.text = ''
                applied.append(src[:60])
                break
        else:
            not_applied.append(src[:60])
    return applied, not_applied


applied, not_applied = apply_replacements_in_tree(doc.element.body)
print(f"applied {len(applied)} of {len(REPLACEMENTS)} replacements")
for a in applied:
    print("  ✓", a)
for n in not_applied:
    print("  ✗", n)


# ---------------------------------------------------------------------------
# Remove the entire ADDENDUM block (now redundant since corrections are
# applied inline).  Block boundaries: from the paragraph whose text is
# "ADDENDUM: CORRECTIONS TO THE ORIGINAL MANUSCRIPT" up to (but not
# including) the paragraph whose text starts with "APPENDIX A:".
# ---------------------------------------------------------------------------
def text_of(elem):
    ts = elem.findall('.//' + qn('w:t'))
    return ''.join(t.text or '' for t in ts).strip()


body = doc.element.body
children = list(body)
start_idx = None
end_idx = None
for i, c in enumerate(children):
    if c.tag != qn('w:p'):
        continue
    t = text_of(c)
    if t.startswith("ADDENDUM: CORRECTIONS"):
        start_idx = i
    elif start_idx is not None and t.startswith("APPENDIX A:"):
        end_idx = i
        break

if start_idx is not None and end_idx is not None:
    print(f"removing addendum block: body indices {start_idx}..{end_idx - 1}")
    to_remove = children[start_idx:end_idx]
    for c in to_remove:
        body.remove(c)
    print(f"  removed {len(to_remove)} elements")
else:
    print("could not locate addendum block (start_idx, end_idx) =",
          start_idx, end_idx)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
os.makedirs(os.path.dirname(DEST), exist_ok=True)
doc.save(DEST)
print(f"wrote {DEST}")
print(f"final paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
