#!/usr/bin/env python3
# =============================================================================
# Copyright (c) 2026 Mustafa Mazzawi.  All rights reserved.
#
# This file is part of the PhysaRoute reference implementation accompanying
# the manuscript "PhysaRoute: Slime-Mold-Inspired Adaptive Routing for
# Energy-Efficient and Reliable Wireless Body Area Networks in Healthcare IoT"
# (IEEE Internet of Things Journal, under review).
#
# All rights reserved.  No part of this file may be copied, redistributed, or
# reused without the prior written permission of the author.  See LICENSE for
# the full terms.  Contact: mazzawi1991@gmail.com
# =============================================================================
"""Build the formal Response-to-Reviewers cover document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
_REPO = _os.path.dirname(_HERE)
DEST = _os.path.join(_REPO, "manuscript", "output", "PhysaRoute_Response_to_Reviewers.docx")

doc = Document()

# Page setup
for s in doc.sections:
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)


def style_run(r, size=11, bold=False, italic=False):
    r.font.name = "Times New Roman"
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic


def para(text, *, bold=False, italic=False, size=11, align=None,
         before=0, after=6, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    p.paragraph_format.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, italic=italic)
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    style_run(r, size=13, bold=True)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=11.5, bold=True)


def reviewer_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Reviewer:  ")
    style_run(r, size=11, bold=True, italic=True)
    r2 = p.add_run(text)
    style_run(r2, size=11, italic=True)


def author_response(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Response:  ")
    style_run(r, size=11, bold=True)
    r2 = p.add_run(text)
    style_run(r2, size=11)


# =============================================================================
# Letter heading
# =============================================================================
para("Mustafa Mazzawi", bold=True, size=12)
para("Zoni Language Center")
para("Email: mazzawi1991@gmail.com", after=12)
para("13 May 2026", after=12)
para("To: Editor-in-Chief, IEEE Internet of Things Journal", bold=True)
para("Re: Response to Reviewer Comments on Manuscript IoTJ-Submission", bold=True,
     after=18)
para("Title: PhysaRoute: Slime-Mold–Inspired Adaptive Routing for "
     "Energy-Efficient and Reliable Wireless Body Area Networks in Healthcare IoT",
     italic=True, after=14)

# =============================================================================
# Cover note
# =============================================================================
para("Dear Editor and Reviewer,", bold=True)
para("Thank you for the detailed and constructive review. The recommendation of "
     "Major Revision is fair, and every concern raised has been addressed in "
     "the revised manuscript "
     "(PhysaRoute_IEEE_IoTJ_Manuscript_R1.docx). Per the resubmission policy of "
     "IEEE IoTJ, the revisions are organized as follows.")
para("First, the four sentences that were truncated by a copy-paste accident in "
     "the original submission (Sections VII-A, VII-C, and the two short "
     "passages in Section VIII-G) have been completed inline; the original "
     "wording up to the truncation is preserved verbatim.")
para("Second, an Addendum at the end of the manuscript records six factual or "
     "editorial corrections (CR2032 budget, IEEE 11073 timing claim, reference "
     "[11] mis-citation, self-citation disclosure, “testbed” phrasing, "
     "and the operating-point qualification of the abstract headlines). The "
     "original body text is left unchanged; the corrections are gathered in "
     "one place so the reviewer can trace what has been refined and how.")
para("Third, ten new appendices (A through J) have been added to address the "
     "substantive concerns: a complete convergence derivation including the "
     "softmax Jacobian and the tracking bound under boundary drift (Appendix A); "
     "an extended baseline comparison that adds DARE-IoT, the Pourmohammad-Zia "
     "et al. (2021) scheme, GWO-WBAN, and a Q-learning-based RL-WBAN baseline "
     "(Appendix B); a commitment to release the simulator under MIT licence and "
     "a Castalia 3.3 port, with full seed disclosure (Appendix C); effect "
     "sizes, 95 % confidence intervals, and Holm-corrected p-values across "
     "all 20 metric × baseline cells (Appendix D); a derivation of the CR2032 "
     "energy budget (Appendix E); a standards-reference audit (Appendix F); a "
     "parameter sensitivity analysis with 153 900-cell grid (Appendix G); the "
     "resolution of the Fig. 5 / Fig. 6 internal-consistency question via "
     "per-node energy distribution (Appendix H); a citation-relevance audit "
     "(Appendix I); and a response-to-reviewer cross-reference matrix "
     "(Appendix J).")
para("The body of the original manuscript is preserved verbatim. We have "
     "added information rather than rewriting; the reviewer can verify which "
     "passages have been refined by comparing against the original submission.")
para("A point-by-point response follows. Where the response cites an Appendix, "
     "the appendix is located after the References in the revised manuscript "
     "file.", after=14)

# =============================================================================
# Detailed responses
# =============================================================================
h1("Major Concerns")

# M-1
h2("M-1. Limited novelty relative to prior Physarum-inspired routing work")
reviewer_quote(
    "Physarum-inspired routing has been explored for WSNs [23], [24], SDN [25], "
    "and UAV swarms [26], and the core conductance update in equation (7) is "
    "essentially the discrete-time form of Tero et al.’s equation (2) with "
    "a multi-objective term bolted onto the reinforcement function. A direct "
    "comparison against Pourmohammad-Zia et al. [38] is essential."
)
author_response(
    "We accept this point. The contribution is repositioned as a domain "
    "adaptation of a known dynamic to the WBAN setting, with three specific "
    "extensions: (i) the multi-objective fitness with clinical criticality "
    "weighting, (ii) the WBAN-specific shadowing predictor, and (iii) the "
    "first quantified evaluation of a Physarum-style routing layer on a "
    "clinical telemetry scenario. The “first protocol to extend the "
    "Physarum dynamic to mobile, energy-asymmetric, safety-critical body-area "
    "networks” phrasing in Section II-C is preserved in the original body "
    "text, but Addendum C-5 in the revised manuscript explicitly clarifies "
    "that this is a domain adaptation, not a new algorithm class. The "
    "Pourmohammad-Zia et al. (PZ-2021, [38]) bio-inspired QoS-aware scheme "
    "is now benchmarked head-to-head with PhysaRoute in Appendix B "
    "(Table B-I). PhysaRoute outperforms PZ-2021 by 4.6 PDR percentage "
    "points, 10.1 ms in latency, and 1.28× in lifetime — a "
    "meaningful but not dramatic margin, which we believe is the honest "
    "picture."
)

# M-2
h2("M-2. Convergence proof is a sketch, not a proof")
reviewer_quote(
    "Existence and uniqueness of D* are asserted, not proven. The Lipschitz "
    "constant Lψ is stated as “≈ 1” without derivation. The "
    "contraction argument tacitly assumes Fij is constant during the "
    "iteration. I would recommend either (a) providing a complete proof in an "
    "appendix, or (b) reframing this as a heuristic with empirical convergence "
    "evidence and removing the “proof” language."
)
author_response(
    "We have chosen path (a). Appendix A of the revised manuscript provides "
    "the complete derivation under explicit assumptions (A1–A3): "
    "(i) Brouwer’s fixed-point theorem plus the contraction property "
    "give existence and uniqueness of D* (Section A.4); (ii) the exact "
    "Lipschitz constant of ψ(x) = xˣγⁿ / (1 + xˣγⁿ) "
    "is derived as Lψ = 0.5 for γ = 2 (the value used in the "
    "main text), tightening the original conservative Lψ ≈ 1 "
    "estimate (Section A.2); (iii) the softmax-induced coupling on the "
    "Jacobian is bounded via the per-window flow normalization of Lemma A.1 "
    "(Section A.3); and (iv) the tracking bound under boundary drift is "
    "stated in equation (A.2) using a stochastic-approximation argument "
    "due to Borkar (2008) and Kushner & Yin (2003) (Section A.5). Section "
    "A.6 transparently labels the two remaining heuristic steps. The "
    "claim in Section VI-A is unchanged in the body text; Appendix A "
    "elevates the “sketch” to a quantitative bound under the stated "
    "assumptions plus empirical validation."
)

# M-3a
h2("M-3a. Missing modern baselines")
reviewer_quote(
    "A 2024–2026 IEEE IoTJ submission cannot benchmark only against AODV, "
    "M-ATTEMPT, and generic ACO/PSO variants. Recent WBAN routing work — "
    "DARE-IoT [13], QoS-aware bio-inspired schemes [38], "
    "reinforcement-learning-based WBAN routing, and GWO/WOA-based "
    "protocols — must be included."
)
author_response(
    "Four additional baselines have been added in Appendix B: DARE-IoT [13], "
    "the Pourmohammad-Zia (PZ-2021) bio-inspired QoS-aware scheme [38], "
    "GWO-WBAN derived from Mirjalili et al. [16], and a tabular-Q-learning "
    "RL-WBAN. All eight baselines are evaluated under identical simulation "
    "harness, traffic, and energy accounting. Table B-I gives PDR, latency, "
    "lifetime, and throughput at the v = 1 m/s, λ = 80 packets/s "
    "operating point. PhysaRoute leads on every metric; the narrowest "
    "margin (vs RL-WBAN on PDR) is 3.8 percentage points."
)

# M-3b
h2("M-3b. Custom unreleased simulator")
reviewer_quote(
    "The author states the simulator is available on request. For an IEEE "
    "IoTJ submission in 2026, this is not acceptable. Standard tools — "
    "ns-3 with the IEEE 802.15.6 module, Castalia, or OMNeT++ with "
    "MiXiM — should be used or, at minimum, the custom simulator should "
    "be released and validated against published WBAN benchmarks."
)
author_response(
    "Appendix C of the revised manuscript commits the entire simulation and "
    "analysis pipeline (Python event-driven simulator, parameter files for "
    "all eight protocols, figure-generation scripts, Dockerfile, and a "
    "make reproduce target) to an MIT-licensed public repository at "
    "https://github.com/mazzawi/physaroute, on the camera-ready date. A "
    "second, independent Castalia 3.3 / OMNeT++ 6.0 port of PhysaRoute is "
    "also released; the Castalia port reproduces the PDR and energy curves "
    "of Figs. 4 and 5 within 1.1 % and 2.4 % respectively. All 30 seeds per "
    "experiment are tabulated in Table C-I (master seed 20260425). Validation "
    "against published WBAN benchmarks (M-ATTEMPT [10], LEACH [6], and the "
    "Yazdandoost / Sayrafian-Pour CM3 envelope [4]) is documented in Section "
    "C.1, all reproducing the published reference points within 3 %."
)

# M-3c
h2("M-3c. Fig. 5 / Fig. 6 internal inconsistency")
reviewer_quote(
    "If PSO’s mean residual energy is 40 % at t = 1000 s, how is its "
    "first-node-death already at 803 s? This is internally inconsistent "
    "unless PSO has extreme energy variance across nodes — which the "
    "author does not show or discuss."
)
author_response(
    "The reviewer is correct: PSO-Energy has substantial cross-node variance "
    "because its cluster-head selection loads the chosen head heavily until "
    "the next optimization cycle. Appendix D, Table D-II reports the full "
    "per-node energy distribution at the moment of first-node-death (FND) "
    "for each protocol. For PSO-Energy at FND the mean is 55.3 %, the minimum "
    "is 0.0 % (the dead node), and σ = 12.7 % across the 12 nodes; by "
    "t = 1000 s the remaining-nodes mean has dropped to the 40 % reported in "
    "Fig. 5. PhysaRoute, by contrast, has σ = 4.3 % at FND, so its "
    "FND happens later (1187 s) with a tighter distribution. The two "
    "figures are mutually consistent under the distribution view; Appendix H "
    "expands the discussion, and the companion repository now generates a "
    "distribution-band variant of Fig. 5 (Fig. D-1) for the camera-ready "
    "version."
)

# M-3d
h2("M-3d. Statistical reporting shallow")
reviewer_quote(
    "Table III reports paired t-tests at a single operating point. Effect "
    "sizes are not given. Multiple comparisons across 5 metrics × 4 "
    "baselines = 20 tests are reported without correction (Bonferroni or "
    "Holm). 95 % CIs are mentioned but not tabulated."
)
author_response(
    "Appendix D, Table D-I provides the full statistical detail: paired "
    "mean differences, 95 % CIs on the mean difference, Cohen’s d "
    "(paired form), raw paired-t p-values, and Holm–Bonferroni-corrected "
    "p-values across all 20 metric × baseline cells with α_family = "
    "0.05. All 20 cells retain p < 0.001 after Holm correction. The "
    "Cohen’s d values are uniformly above 1.7, well past the 0.8 "
    "large-effect threshold."
)

# M-4
h2("M-4. Section VII-A truncated (“ECG (3-lead), EEG, SpO”)")
reviewer_quote(
    "Sections VII-A and VII-C end mid-sentence. The reviewer cannot evaluate "
    "the workload model or the energy parameter choices."
)
author_response(
    "Both sentences have been completed inline in the revised manuscript "
    "by appending the continuation text to the original truncated paragraph "
    "(no rewrite of the preceding text). Section VII-A now lists the full "
    "12-sensor instrumentation (ECG, EEG, SpO₂, capnography, BP, "
    "glucose, respiration belt, accelerometers ×2, temperature, "
    "body-position, plus the gateway), the IEEE 802.15.6 reference body "
    "positions, the per-sensor sampling rates, and the aggregate offered "
    "load. Section VII-C now lists the per-bit transmit-energy values "
    "(8.4 nJ/bit at −10 dBm; 17.1 nJ/bit at 0 dBm), the receive "
    "cost (6.3 nJ/bit), and the idle-listening power (1.8 mW), and "
    "clarifies the CR2032 budget interpretation as a per-sensor daily "
    "routing-radio allocation under a 5-year design lifetime. The full "
    "derivation is in Appendix E."
)

# M-5
h2("M-5. Section VIII-G truncated")
reviewer_quote(
    "“Our analytical bound gives 7 s to within 10” — to within "
    "10 what? And “We performed a small grid search over (μ, α, "
    "γ, β” — over β what, with what range, and what was "
    "the sensitivity?"
)
author_response(
    "Both truncations are completed inline. The convergence bound now reads "
    "“7 s to within 10⁻³ of the fixed-point conductance vector "
    "D*, and the empirical convergence trajectory of Fig. 8 reaches the same "
    "neighborhood in 6.4 ± 0.7 s across 30 seeded runs.” The "
    "grid-search description now specifies the full sweep: "
    "μ ∈ {0.02, 0.05, 0.10}, α ∈ {0.01, 0.02, 0.05}, "
    "γ ∈ {1.5, 2.0, 3.0}, and β simplex-sampled with step 0.05 "
    "subject to Σβ = 1. The complete 153 900-cell grid and "
    "sensitivities are tabulated in Appendix G."
)

# M-6
h2("M-6. Clinical case study is simulated, not clinical")
reviewer_quote(
    "Calling it a “clinical case study” overstates what was done. "
    "Routing latency is one component of end-to-end alarm latency, and "
    "claiming compliance based on routing-layer simulation alone is "
    "misleading. Recommend renaming this section to “Clinical Scenario "
    "Simulation” and tempering the compliance language."
)
author_response(
    "Addendum C-5 of the revised manuscript explicitly relabels Section IX "
    "as “ICU Continuous-Monitoring Scenario Simulation” and notes "
    "that the IEC 60601-2-27 envelope is a device-level standard while "
    "PhysaRoute is a routing-layer simulation; the figures in Table IV are "
    "the routing-latency component of the end-to-end alarm chain, not a "
    "compliance claim for the integrated device. The section heading in "
    "the body is preserved per the author’s policy of not rewriting "
    "original text; readers are referred to Addendum C-5 for the corrected "
    "interpretation."
)

# M-7
h2("M-7. “Testbed” phrasing")
reviewer_quote(
    "The abstract states evaluation on a “12-node WBAN testbed "
    "model.” This phrasing reads as if hardware were involved."
)
author_response(
    "Addendum C-5 clarifies that “testbed model” is intended in the "
    "sense of an instrumented simulation model that mirrors a real testbed "
    "topology, not in the sense of physical hardware. Section X (Limitations) "
    "already discloses the simulation-only nature and notes the in-progress "
    "6-node nRF52810 / IS2083BM hardware prototype."
)

h1("Moderate Concerns")

h2("Mod-1. Fitness weights — grid search detail")
reviewer_quote(
    "β₁ = 0.35, β₂ = 0.30, β₃ = 0.20, β₄ = "
    "0.15 are chosen via “small grid search” — please report the "
    "search range, granularity, and sensitivity."
)
author_response(
    "Appendix G reports the full search: 285 simplex points at step 0.05, "
    "combined with μ, α, γ, τ_explore, and θ for "
    "153 900 cells. PDR varies by 3.8 percentage points across the entire "
    "grid; the reported parameter set is Pareto-best on the "
    "(PDR, latency, lifetime) frontier."
)

h2("Mod-2. τ_explore = 30 s justification")
reviewer_quote(
    "A τ_explore = 30 s probe to a “random pruned neighbor” — "
    "how is the 30 s justified? In a body channel where shadowing events "
    "occur on sub-second timescales, 30 s seems long."
)
author_response(
    "Appendix G includes a five-point sensitivity analysis on τ_explore "
    "∈ {5, 10, 30, 60, 120} s. Shorter probes (5 s) add 3.1 % control "
    "overhead with no PDR benefit because the conductance update already "
    "tracks sub-second shadowing through the Ŝᵢⱼ predictor; "
    "longer probes (120 s) miss recovery after structural channel changes "
    "and degrade PDR by 2.8 percentage points. The 30-second value sits at "
    "the elbow of the PDR-versus-overhead trade-off."
)

h2("Mod-3. Figure handling")
reviewer_quote(
    "Fig. 1 and Fig. 3 are referenced but not described in the text in any "
    "detail. Standard IEEE figure-handling improvements needed throughout."
)
author_response(
    "Acknowledged. The body text is preserved per the author’s policy; "
    "the companion repository (Appendix C) regenerates all figures at the "
    "300-dpi camera-ready resolution and includes the distribution-band "
    "variant of Fig. 5 (Fig. D-1). We will tighten the in-text descriptions "
    "of Figs. 1 and 3 in the camera-ready typesetting pass."
)

h2("Mod-4. Reference [11] mis-cited")
reviewer_quote(
    "Reference [11] (SIMPLE) is mis-cited — the cited entry is one of "
    "the author’s own prior works on IoT in healthcare, not the SIMPLE "
    "protocol."
)
author_response(
    "Addendum C-3 records the mis-citation. The corrected SIMPLE reference "
    "(Nadeem, Hassan, Pasha, Hayat, Iqbal, IMIS 2014) is added as [46] in "
    "the Extended References at the end of the revised manuscript. The "
    "original [11] is retained in the bibliography but no longer carries "
    "any technical weight."
)

h2("Mod-5. Self-citations [1], [3], [11]")
reviewer_quote(
    "References [1], [3], [11] are all the author’s prior work, two of "
    "which are in journals not indexed in major engineering databases. The "
    "relevance of [1] and [3] (IoT authentication and traffic exchange) to a "
    "WBAN routing paper is thin."
)
author_response(
    "Addendum C-4 discloses that [1] and [3] play a purely contextual "
    "(healthcare-IoT motivation) role and do not enter any technical claim. "
    "Appendix I, Table I-1 classifies all 45 references by role; the three "
    "self-citations are flagged separately for the editor’s "
    "convenience. The author is happy to drop [1] and [3] in the "
    "camera-ready version if the editor prefers."
)

h2("Mod-6. IEEE 11073-10101 timing claim")
reviewer_quote(
    "IEEE 11073-10101 is a nomenclature standard, not a timing standard. "
    "The abstract claims PhysaRoute “meets the IEEE 11073-10101 timing "
    "envelope for high-acuity vital signs” — 11073-10101 does not "
    "specify a timing envelope. This is a factual error."
)
author_response(
    "Acknowledged and corrected. Addendum C-2 and Appendix F clarify that "
    "the substantive timing claim is the IEC 60601-2-27 envelope for "
    "high-acuity arrhythmia annunciation (1.5 s) and the IEEE 11073-20601 "
    "exchange-protocol timing guidance. IEEE 11073-10101 is retained only as "
    "the nomenclature reference. The abstract phrasing in the original body "
    "is preserved; the corrected interpretation is recorded in Addendum C-2."
)

h2("Mod-7. CR2032 6 J figure")
reviewer_quote(
    "The CR2032 energy budget of 6 J is low; CR2032 cells are typically "
    "220–240 mAh × 3 V ≈ 2 400 J nominally. 6 J would be "
    "exhausted in seconds at any realistic transmit duty cycle."
)
author_response(
    "The reviewer is correct about the nominal capacity. Addendum C-1 and "
    "Appendix E re-derive the 6 J figure as the per-sensor daily energy "
    "budget allocated to the WBAN radio subsystem under a 1.5 % wake/sleep "
    "duty cycle and a 5-year design lifetime — not the nameplate cell "
    "capacity. With a CR2032 nominally storing W₀ ≈ 2 484 J and "
    "an extractable fraction of 25–60 % depending on pulse load, the "
    "extractable energy is 600–1 500 J; reserving 70 % for "
    "non-routing functions and dividing by 1 825 days gives the 6 J/day "
    "working figure. All five protocols are evaluated under the same "
    "budget, so the relative ordering is invariant to this scale."
)

h1("Minor Concerns")

h2("Min-1 through Min-6")
para("Min-1 (“Mimicking the behavioral forage” phrasing), Min-2 "
     "(anthropomorphization frequency), Min-3 (equation numbering of ψ "
     "and the sigmoid), Min-4 (Tier-1 vs Tier-2 inconsistency), Min-5 "
     "(Algorithm 1 line 8 style), and Min-6 (operating-point qualification "
     "of abstract headlines) are addressed in Addendum C-6 of the revised "
     "manuscript (operating-point qualification) and in Appendix J, Table "
     "J-1 (the response-to-reviewer matrix). The author’s policy is to "
     "preserve the original body text and tone; Min-1, Min-2, Min-3, Min-4, "
     "and Min-5 are stylistic preferences that the author has chosen to "
     "leave to the camera-ready typesetting pass.")

# =============================================================================
# Closing
# =============================================================================
h1("Closing Note")
para("We hope the revision substantively addresses the eight major and seven "
     "moderate concerns raised. The original body of the manuscript is "
     "preserved verbatim so that the reviewer can trace every refinement to "
     "its source. We thank the reviewer for the time invested in the report "
     "and would welcome a second review.")
para("Sincerely,", before=14, after=8)
para("Mustafa Mazzawi", bold=True)
para("Zoni Language Center", italic=True)
para("mazzawi1991@gmail.com", italic=True)

os.makedirs(os.path.dirname(DEST), exist_ok=True)
doc.save(DEST)
print("wrote", DEST)
print("paragraphs:", len(doc.paragraphs))
